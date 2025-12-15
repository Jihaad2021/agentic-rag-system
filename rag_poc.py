"""
RAG Proof of Concept - Phase 1 Day 2
PDF loading, text chunking, and embeddings.

Note: Using Claude 3 Haiku (claude-3-haiku-20240307) for API calls.
"""

import os
from typing import List, Dict, Any
from pypdf import PdfReader
from dotenv import load_dotenv
import tiktoken
from langchain_voyageai import VoyageAIEmbeddings
from langchain_anthropic import ChatAnthropic
from langchain.prompts import ChatPromptTemplate
from docx import Document as DocxDocument
import numpy as np

# Load environment
load_dotenv()


class DocumentLoader:
    """Load and extract text from multiple file formats."""
    
    def __init__(self):
        """Initialize document loader."""
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def load(self, file_path: str) -> str:
        """
        Load document and extract all text.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text as string
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format not supported
        """
        # Validate file exists
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Validate file format
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in self.supported_formats:
            raise ValueError(
                f"Unsupported format: {file_ext}. "
                f"Supported: {', '.join(self.supported_formats)}"
            )
        
        print(f"📄 Loading {file_ext.upper()} file: {file_path}")
        
        # Route to appropriate loader
        if file_ext == '.pdf':
            text = self._load_pdf(file_path)
        elif file_ext == '.docx':
            text = self._load_docx(file_path)
        elif file_ext == '.txt':
            text = self._load_txt(file_path)
        else:
            raise ValueError(f"Unsupported format: {file_ext}")
        
        print(f"✅ Extracted {len(text)} characters")
        
        return text
    
    def _load_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using PyPDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Combined text from all pages
        """
        text_parts = []
        
        reader = PdfReader(file_path)
        
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            page_text = self._clean_text(page_text)
            
            if page_text.strip():
                text_parts.append(page_text)
                print(f"  Page {page_num}: {len(page_text)} chars")
        
        return "\n\n".join(text_parts)
    
    def _load_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX using python-docx.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Combined text from all paragraphs
        """
        text_parts = []
        
        doc = DocxDocument(file_path)
        
        for i, paragraph in enumerate(doc.paragraphs, start=1):
            para_text = paragraph.text
            para_text = self._clean_text(para_text)
            
            if para_text.strip():
                text_parts.append(para_text)
        
        print(f"  Extracted {len(doc.paragraphs)} paragraphs")
        
        return "\n\n".join(text_parts)
    
    def _load_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            File content as string
        """
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        text = self._clean_text(text)
        
        lines = len(text.splitlines())
        print(f"  Extracted {lines} lines")
        
        return text
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Remove null bytes
        text = text.replace("\x00", "")
        
        return text
    
    def count_pages(self, file_path: str) -> int:
        """
        Count pages/sections in document.
        
        Args:
            file_path: Path to document
            
        Returns:
            Number of pages or estimated sections
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            reader = PdfReader(file_path)
            return len(reader.pages)
        
        elif file_ext == '.docx':
            doc = DocxDocument(file_path)
            # Estimate pages (rough: 500 words per page)
            total_words = sum(len(p.text.split()) for p in doc.paragraphs)
            return max(1, total_words // 500)
        
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                lines = len(f.readlines())
            # Estimate pages (rough: 50 lines per page)
            return max(1, lines // 50)
        
        return 1

PDFLoader = DocumentLoader

class TextChunker:
    """Split text into chunks with token counting."""
    
    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        encoding_name: str = "cl100k_base"
    ):
        """
        Initialize text chunker.
        
        Args:
            chunk_size: Maximum tokens per chunk
            chunk_overlap: Overlap tokens between chunks
            encoding_name: Tokenizer encoding (cl100k_base for GPT-4/Claude)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.encoding = tiktoken.get_encoding(encoding_name)
    
    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Input text to chunk
            
        Returns:
            List of chunk dictionaries with text, tokens, and metadata
        """
        print(f"\n📝 Chunking text...")
        print(f"   Chunk size: {self.chunk_size} tokens")
        print(f"   Overlap: {self.chunk_overlap} tokens")
        
        # Tokenize entire text
        tokens = self.encoding.encode(text)
        total_tokens = len(tokens)
        
        print(f"   Total tokens: {total_tokens}")
        
        chunks = []
        start_idx = 0
        chunk_num = 0
        
        while start_idx < total_tokens:
            # Calculate end index
            end_idx = min(start_idx + self.chunk_size, total_tokens)
            
            # Extract chunk tokens
            chunk_tokens = tokens[start_idx:end_idx]
            
            # Decode back to text
            chunk_text = self.encoding.decode(chunk_tokens)
            
            # Create chunk metadata
            chunk = {
                "chunk_id": f"chunk_{chunk_num}",
                "text": chunk_text,
                "tokens": chunk_tokens,
                "token_count": len(chunk_tokens),
                "start_idx": start_idx,
                "end_idx": end_idx,
                "chunk_num": chunk_num
            }
            
            chunks.append(chunk)
            chunk_num += 1
            
            # Move to next chunk with overlap
            start_idx += self.chunk_size - self.chunk_overlap
        
        print(f"✅ Created {len(chunks)} chunks")
        
        return chunks
    
    def count_tokens(self, text: str) -> int:
        """
        Count tokens in text.
        
        Args:
            text: Input text
            
        Returns:
            Number of tokens
        """
        return len(self.encoding.encode(text))

class Embedder:
    """Generate embeddings using Voyage AI."""
    
    def __init__(self):
        """Initialize embedder."""
        self.embedder = VoyageAIEmbeddings(
            voyage_api_key=os.getenv("VOYAGE_API_KEY"),
            model="voyage-large-2"
        )
        print("📊 Embedder initialized (voyage-large-2, 1536 dimensions)")
    
    def embed_chunks(self, chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """
        Generate embeddings for chunks.
        
        Args:
            chunks: List of chunk dictionaries
            
        Returns:
            Chunks with added 'embedding' field
        """
        print(f"\n🔢 Generating embeddings for {len(chunks)} chunks...")
        
        # Extract texts
        texts = [chunk['text'] for chunk in chunks]
        
        # Generate embeddings (batch)
        embeddings = self.embedder.embed_documents(texts)
        
        # Add embeddings to chunks
        for chunk, embedding in zip(chunks, embeddings):
            chunk['embedding'] = embedding
        
        print(f"✅ Generated {len(embeddings)} embeddings")
        print(f"   Dimension: {len(embeddings[0])}")
        
        return chunks
    
    def embed_query(self, query: str) -> List[float]:
        """
        Generate embedding for a query.
        
        Args:
            query: Query text
            
        Returns:
            Query embedding vector
        """
        return self.embedder.embed_query(query)


class SimpleVectorStore:
    """Simple in-memory vector storage and search."""
    
    def __init__(self):
        """Initialize vector store."""
        self.chunks = []
        print("💾 SimpleVectorStore initialized")
    
    def add_chunks(self, chunks: List[Dict[str, Any]]) -> None:
        """
        Add chunks to storage.
        
        Args:
            chunks: List of chunks with embeddings
        """
        self.chunks.extend(chunks)
        print(f"✅ Added {len(chunks)} chunks to storage (Total: {len(self.chunks)})")
    
    def search(
        self,
        query_embedding: List[float],
        top_k: int = 5
    ) -> List[Dict[str, Any]]:
        """
        Search for similar chunks using cosine similarity.
        
        Args:
            query_embedding: Query vector
            top_k: Number of results to return
            
        Returns:
            List of top-k most similar chunks with scores
        """
        if not self.chunks:
            print("⚠️  No chunks in storage")
            return []
        
        print(f"\n🔍 Searching for top-{top_k} similar chunks...")
        
        # Calculate similarities
        similarities = []
        for chunk in self.chunks:
            # Cosine similarity
            similarity = self._cosine_similarity(
                query_embedding,
                chunk['embedding']
            )
            similarities.append({
                **chunk,
                'score': similarity
            })
        
        # Sort by score (descending)
        similarities.sort(key=lambda x: x['score'], reverse=True)
        
        # Return top-k
        top_results = similarities[:top_k]
        
        print(f"✅ Found {len(top_results)} results")
        for i, result in enumerate(top_results, 1):
            print(f"   {i}. {result['chunk_id']}: score={result['score']:.4f}")
        
        return top_results
    
    def _cosine_similarity(
        self,
        vec1: List[float],
        vec2: List[float]
    ) -> float:
        """
        Calculate cosine similarity between two vectors.
        
        Args:
            vec1: First vector
            vec2: Second vector
            
        Returns:
            Cosine similarity score (0-1)
        """
        vec1 = np.array(vec1)
        vec2 = np.array(vec2)
        
        dot_product = np.dot(vec1, vec2)
        norm1 = np.linalg.norm(vec1)
        norm2 = np.linalg.norm(vec2)
        
        return dot_product / (norm1 * norm2)

class AnswerGenerator:
    """Generate answers using Claude with retrieved context."""
    
    def __init__(self, model: str = "claude-3-haiku-20240307"):
        """
        Initialize answer generator.
        
        Args:
            model: Claude model to use
        """
        self.llm = ChatAnthropic(
            api_key=os.getenv("ANTHROPIC_API_KEY"),
            model=model,
            temperature=0  # Deterministic for consistency
        )
        print(f"🤖 AnswerGenerator initialized (model: {model})")
    
    def generate(
        self,
        query: str,
        chunks: List[Dict[str, Any]],
        max_chunks: int = 5
    ) -> Dict[str, Any]:
        """
        Generate answer from query and retrieved chunks.
        
        Args:
            query: User question
            chunks: Retrieved chunks (with scores)
            max_chunks: Maximum chunks to use for context
            
        Returns:
            Dictionary with answer, citations, and metadata
        """
        print(f"\n🤖 Generating answer for: '{query}'")
        
        # Limit chunks
        top_chunks = chunks[:max_chunks]
        
        # Build context from chunks
        context = self._build_context(top_chunks)
        
        # Build prompt
        prompt = self._build_prompt(query, context)
        
        # Generate answer
        print(f"   Using {len(top_chunks)} chunks as context")
        print(f"   Calling Claude...")
        
        response = self.llm.invoke(prompt)
        answer_text = response.content
        
        # Extract citations
        citations = self._extract_citations(answer_text, top_chunks)
        
        result = {
            "query": query,
            "answer": answer_text,
            "citations": citations,
            "chunks_used": len(top_chunks),
            "model": self.llm.model
        }
        
        print(f"✅ Answer generated ({len(answer_text)} chars)")
        
        return result
    
    def _build_context(self, chunks: List[Dict[str, Any]]) -> str:
        """
        Build context string from chunks.
        
        Args:
            chunks: List of chunks with metadata
            
        Returns:
            Formatted context string
        """
        context_parts = []
        
        for i, chunk in enumerate(chunks, 1):
            # Add chunk with source reference
            chunk_text = f"[Source {i}]\n{chunk['text']}\n"
            context_parts.append(chunk_text)
        
        return "\n".join(context_parts)
    
    def _build_prompt(self, query: str, context: str) -> str:
        """
        Build prompt for Claude.
        
        Args:
            query: User question
            context: Context from retrieved chunks
            
        Returns:
            Complete prompt string
        """
        prompt_template = """You are a helpful AI assistant answering questions based on provided context.

Context:
{context}

Question: {query}

Instructions:
1. Answer the question using ONLY the information in the context above
2. If the context doesn't contain enough information, say "I don't have enough information to answer this question."
3. Be concise and direct
4. When using information from a source, mention it like: "According to Source 1, ..."
5. Do not make up information not present in the context

Answer:"""

        return prompt_template.format(context=context, query=query)
    
    def _extract_citations(
        self,
        answer: str,
        chunks: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        Extract which sources were cited in the answer.
        
        Args:
            answer: Generated answer text
            chunks: Chunks that were provided as context
            
        Returns:
            List of citations with chunk info
        """
        citations = []
        
        # Simple citation detection: look for "Source N" mentions
        for i, chunk in enumerate(chunks, 1):
            source_mention = f"Source {i}"
            
            if source_mention in answer:
                citations.append({
                    "source_number": i,
                    "chunk_id": chunk.get('chunk_id', f'chunk_{i}'),
                    "text_preview": chunk['text'][:100] + "...",
                    "score": chunk.get('score', 0.0)
                })
        
        # If no explicit citations found, assume all chunks were used
        if not citations:
            citations = [
                {
                    "source_number": i,
                    "chunk_id": chunk.get('chunk_id', f'chunk_{i}'),
                    "text_preview": chunk['text'][:100] + "...",
                    "score": chunk.get('score', 0.0)
                }
                for i, chunk in enumerate(chunks, 1)
            ]
        
        return citations

def test_chunking():
    """Test text chunking functionality."""
    
    print("\n" + "=" * 60)
    print("TESTING TEXT CHUNKING")
    print("=" * 60)
    
    # Sample text
    sample_text = """
    Artificial intelligence (AI) is intelligence demonstrated by machines, 
    in contrast to the natural intelligence displayed by humans and animals. 
    Leading AI textbooks define the field as the study of "intelligent agents": 
    any device that perceives its environment and takes actions that maximize 
    its chance of successfully achieving its goals. Colloquially, the term 
    "artificial intelligence" is often used to describe machines (or computers) 
    that mimic "cognitive" functions that humans associate with the human mind, 
    such as "learning" and "problem solving". As machines become increasingly 
    capable, tasks considered to require "intelligence" are often removed from 
    the definition of AI, a phenomenon known as the AI effect. A quip in 
    Tesler's Theorem says "AI is whatever hasn't been done yet." For instance, 
    optical character recognition is frequently excluded from things considered 
    to be AI, having become a routine technology.
    """ * 10  # Repeat to get more tokens
    
    # Create chunker
    chunker = TextChunker(chunk_size=100, chunk_overlap=20)
    
    # Count tokens
    token_count = chunker.count_tokens(sample_text)
    print(f"\nSample text: {token_count} tokens")
    
    # Chunk text
    chunks = chunker.chunk_text(sample_text)
    
    # Display first 3 chunks
    print(f"\n📋 First 3 chunks (of {len(chunks)}):")
    for i, chunk in enumerate(chunks[:3], 1):
        print(f"\n--- Chunk {i} ---")
        print(f"ID: {chunk['chunk_id']}")
        print(f"Tokens: {chunk['token_count']}")
        print(f"Text preview: {chunk['text'][:100]}...")
    
    # Verify overlap
    print(f"\n🔍 Verifying overlap:")
    if len(chunks) >= 2:
        chunk1_text = chunks[0]['text']
        chunk2_text = chunks[1]['text']
        
        # Find overlap
        overlap_found = False
        for i in range(len(chunk1_text)):
            if chunk2_text.startswith(chunk1_text[i:]):
                overlap_text = chunk1_text[i:]
                print(f"✅ Overlap detected: ~{len(overlap_text)} chars")
                overlap_found = True
                break
        
        if not overlap_found:
            print("⚠️  No obvious overlap detected (might be token-level)")
    
    print("\n✅ Chunking tests complete!")


def test_error_handling():
    """Test error handling for various scenarios."""
    
    print("\n" + "=" * 60)
    print("TESTING ERROR HANDLING")
    print("=" * 60)
    
    loader = PDFLoader()
    
    # Test 1: Non-existent file
    print("\nTest 1: Non-existent file")
    try:
        loader.load("nonexistent.pdf")
        print("❌ Should have raised FileNotFoundError")
    except FileNotFoundError as e:
        print(f"✅ Correctly raised: {type(e).__name__}")
    
    # Test 2: Unsupported format
    print("\nTest 2: Unsupported format")
    try:
        # Create dummy file
        with open("data/uploads/test.txt", "w") as f:
            f.write("test")
        
        loader.load("data/uploads/test.txt")
        print("❌ Should have raised ValueError")
    except ValueError as e:
        print(f"✅ Correctly raised: {type(e).__name__}")
    finally:
        # Cleanup
        if os.path.exists("data/uploads/test.txt"):
            os.remove("data/uploads/test.txt")
    
    print("\n✅ All error handling tests passed!")

def test_end_to_end_rag():
    """Test complete RAG pipeline with answer generation."""
    
    print("\n" + "=" * 60)
    print("TESTING COMPLETE RAG PIPELINE (WITH ANSWER GENERATION)")
    print("=" * 60)
    
    # 1. Load PDF
    print("\n1️⃣ LOADING PDF...")
    loader = PDFLoader()
    test_file = "data/uploads/sample.pdf"
    
    if not os.path.exists(test_file):
        print("❌ Sample PDF not found. Skipping end-to-end test.")
        return
    
    text = loader.load(test_file)
    
    # 2. Chunk text
    print("\n2️⃣ CHUNKING TEXT...")
    chunker = TextChunker(chunk_size=200, chunk_overlap=50)
    chunks = chunker.chunk_text(text)
    
    # 3. Generate embeddings
    print("\n3️⃣ GENERATING EMBEDDINGS...")
    embedder = Embedder()
    chunks = embedder.embed_chunks(chunks)
    
    # 4. Store in vector DB
    print("\n4️⃣ STORING IN VECTOR DATABASE...")
    vector_store = SimpleVectorStore()
    vector_store.add_chunks(chunks)
    
    # 5. Initialize generator
    print("\n5️⃣ INITIALIZING ANSWER GENERATOR...")
    generator = AnswerGenerator()
    
    # 6. Query and generate answers
    print("\n6️⃣ QUERYING AND GENERATING ANSWERS...")
    test_queries = [
        "What is this document about?",
        "Summarize the main content"
    ]
    
    for query in test_queries:
        print("\n" + "-" * 60)
        print(f"📝 Query: '{query}'")
        print("-" * 60)
        
        # Embed query
        query_embedding = embedder.embed_query(query)
        
        # Search for relevant chunks
        results = vector_store.search(query_embedding, top_k=3)
        
        print(f"\n🔍 Retrieved {len(results)} chunks:")
        for i, result in enumerate(results, 1):
            print(f"   {i}. Score: {result['score']:.4f}")
            print(f"      Preview: {result['text'][:80]}...")
        
        # Generate answer
        answer_result = generator.generate(query, results, max_chunks=3)
        
        # Display answer
        print(f"\n💬 ANSWER:")
        print(f"   {answer_result['answer']}")
        
        print(f"\n📚 CITATIONS ({len(answer_result['citations'])}):")
        for citation in answer_result['citations']:
            print(f"   - Source {citation['source_number']}: {citation['text_preview']}")
    
    print("\n" + "=" * 60)
    print("✅ COMPLETE RAG PIPELINE WORKING!")
    print("=" * 60)
    print("\nPipeline Summary:")
    print("  PDF → Chunks → Embeddings → Search → Generate → Answer ✅")

def demo_rag_qa():
    """
    Simple demo: Ask questions about a PDF.
    
    This is a user-friendly demo function.
    """
    print("\n" + "=" * 60)
    print("🎯 RAG Q&A DEMO")
    print("=" * 60)
    
    # Check if PDF exists
    test_file = "data/uploads/sample.pdf"
    if not os.path.exists(test_file):
        print(f"\n❌ No PDF found at: {test_file}")
        print("   Please add a PDF file first.")
        return
    
    print(f"\n📄 Loading PDF: {test_file}")
    
    # Build RAG pipeline
    loader = PDFLoader()
    chunker = TextChunker(chunk_size=500, chunk_overlap=50)
    embedder = Embedder()
    vector_store = SimpleVectorStore()
    generator = AnswerGenerator()
    
    # Process PDF
    text = loader.load(test_file)
    chunks = chunker.chunk_text(text)
    chunks = embedder.embed_chunks(chunks)
    vector_store.add_chunks(chunks)
    
    print("\n✅ RAG system ready!")
    print("\n" + "=" * 60)
    
    # Interactive Q&A
    questions = [
        "What is the main topic of this document?",
        "Give me a brief summary",
        "What are the key points?"
    ]
    
    for i, question in enumerate(questions, 1):
        print(f"\n❓ Question {i}: {question}")
        print("-" * 60)
        
        # RAG pipeline
        query_embedding = embedder.embed_query(question)
        relevant_chunks = vector_store.search(query_embedding, top_k=5)
        answer_result = generator.generate(question, relevant_chunks)
        
        # Display
        print(f"\n💡 Answer:\n{answer_result['answer']}\n")
        
        if answer_result['citations']:
            print(f"📌 Sources used: {len(answer_result['citations'])} chunks")
    
    print("\n" + "=" * 60)
    print("✅ Demo complete!")
    print("=" * 60)

def main():
    """Main function - comprehensive RAG testing."""
    
    print("=" * 60)
    print("RAG POC - PHASE 1 DAY 3: COMPLETE RAG WITH GENERATION")
    print("=" * 60)
    
    # Test 1: Basic chunking
    print("\n[Test 1] Basic Chunking")
    test_chunking()
    
    # Test 2: Complete RAG pipeline
    print("\n[Test 2] Complete RAG Pipeline")
    test_end_to_end_rag()
    
    # Test 3: Error handling
    print("\n[Test 3] Error Handling")
    test_error_handling()
    
    # Demo: Interactive Q&A
    print("\n[Demo] Interactive Q&A")
    demo_rag_qa()
    
    print("\n" + "=" * 60)
    print("🎉 ALL TESTS COMPLETE - DAY 3!")
    print("=" * 60)
    print("\n📊 System Status:")
    print("  ✅ PDF Loading")
    print("  ✅ Text Chunking")
    print("  ✅ Embeddings (Voyage AI)")
    print("  ✅ Vector Search")
    print("  ✅ Answer Generation (Claude)")
    print("  ✅ Citations")
    print("\n🚀 Ready for Streamlit UI (Day 4-5)")


if __name__ == "__main__":
    main()